import re
import docx
from datetime import datetime


class ProtocolParser:

    def __init__(self, name):
        self.doc = docx.Document(name)
        doc_text = '\n'.join([paragraph.text for paragraph in self.doc.paragraphs])
        self.start_paragraph = self.doc.paragraphs[0]
        self.result_text = ''
        for paragraph in self.doc.paragraphs:
            if "Решение:" in paragraph.text:
                self.start_paragraph = paragraph
        if self.start_paragraph is not None:
            self.result_text = self.start_paragraph.text

    def get_list_of_dates(self):
        date_regex = r"\d{2}\.\d{2}\.\d{4}"
        date_list = re.findall(date_regex, self.result_text)
        return date_list

    def get_list_of_times(self):
        time_regex = r"\d{2} час. \d{2} мин."
        time_list = re.findall(time_regex, self.result_text)
        time_objs = [datetime.strptime(t, '%H час. %M мин.') for t in time_list]
        formatted_times = [t.strftime('%H:%M') for t in time_objs]
        return formatted_times

    def get_form(self):
        form_regex = r"очной|заочной|очно-заочной"
        form_match = re.search(form_regex, self.result_text)
        form = form_match.group(0)
        return form

    def get_type(self):
        type_regex = r"очередное|внеочередное"
        type_match = re.search(type_regex, self.result_text)
        type = type_match.group(0)
        return type

    def get_questions(self):
        need_to_write = False
        questions = []
        for paragraph in self.doc.paragraphs:
            if "Решение:" in paragraph.text:
                need_to_write = True
            elif "Приложение:" in paragraph.text:
                need_to_write = False
            elif need_to_write:
                questions.append(paragraph.text)
        return questions

    def format_datetime(self, form, dates, times):
        if form == 'очной':  # 2 1
            return f'{times[0]}-{times[1]} {dates[0]}'
        elif form == 'заочной':  # 2 2
            return f'{times[0]}-{times[1]} {dates[0]}-{dates[1]}'
        else:  # 4 3
            return f'{times[0]}-{times[1]} {dates[0]} / {times[2]}-{times[3]} {dates[1]}-{dates[2]}'

    def get_place(self, form):
        if form == 'очной':  # 2 1
            return 'Ул. Ленина, д.14'
        elif form == 'заочной':  # 2 2
            return 'Сайт СНТ'
        else:  # 4 3
            return 'Ул. Ленина, д.14 / Сайт СНТ'
