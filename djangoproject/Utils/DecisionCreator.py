import docx
import time as t
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


def create_decision(name, type, date_time, form, questions, zaoch_list, protocol_date, current_date):
    new_form = form.replace('ой', 'ого')
    type = type.replace('ое', 'ого')
    total_ochno = questions[0].yes + questions[0].no + questions[0].idk
    total_zaochno = zaoch_list[0][0]+zaoch_list[0][1]+zaoch_list[0][2]

    if (total_ochno + total_zaochno) > 0.5*49:
        quorum = 'да'
    else:
        quorum = 'нет'

    doc = docx.Document(name)
    for para in doc.paragraphs:
        text = para.text
        text = text.replace('<<form>>', new_form)
        text = text.replace('<<type>>', type)
        text = text.replace('<<current_date>>', current_date)
        para.text = text

    for question in questions:
        p = doc.add_paragraph(question.question)
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = doc.add_paragraph(f'Решение: {question.decision}')
        p.paragraph_format.first_line_indent = Mm(7.5)

    if new_form == 'очно-заочного':
        doc.add_paragraph()
        p = doc.add_paragraph(
            'Решение принято на основании голосования в очно-заочной форме путем суммы голосов за очную часть голосования и заочную часть голосования.')
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = doc.add_paragraph(f'Общее количество членов СНТ 1: 49')
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = doc.add_paragraph(
            f'Очно: {total_ochno}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = doc.add_paragraph(
            f'Заочно: {total_zaochno}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        if quorum == 'да':
            p = doc.add_paragraph(
                f'Кворум установлен ({total_ochno + total_zaochno} голосов в сумме)'
            )
            p.paragraph_format.first_line_indent = Mm(7.5)
        else:
            p = doc.add_paragraph(
                f'Кворум не установлен ({total_ochno + total_zaochno} голосов в сумме)'
            )
            p.paragraph_format.first_line_indent = Mm(7.5)

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph(f'Председатель правления СНТ « СНТ 1».')
    p.paragraph_format.first_line_indent = Mm(7.5)
    p = doc.add_paragraph('[подпись]_____________/[расшифровка] Кузнецов С.М.')
    p.paragraph_format.first_line_indent = Mm(7.5)
    p = doc.add_paragraph(f'Дата {current_date}')
    p.paragraph_format.first_line_indent = Mm(7.5)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    for para in range(4, len(doc.paragraphs)):
        doc.paragraphs[para].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        doc.paragraphs[para].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    name = f'reshenie_{int(t.time())}.docx'
    doc.save(f'./uploads/{name}')
    return name