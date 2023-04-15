import re
import docx
from datetime import datetime


class ProtocolParser:

    def __init__(self, name):
        self.doc = docx.Document(name)
        doc_text = '\n'.join([paragraph.text for paragraph in self.doc.paragraphs])
        self.start_paragraph = self.doc.paragraphs[0]
        self.result_text = ''
        for paragraph in self.doc.paragraphs:
            if "Решение:" in paragraph.text:
                self.start_paragraph = paragraph
        if self.start_paragraph is not None:
            self.result_text = self.start_paragraph.text

    def get_list_of_dates(self):
        date_regex = r"\d{2}\.\d{2}\.\d{4}"
        return re.findall(date_regex, self.result_text)

    def get_list_of_times(self):
        time_regex = r"\d{2} час. \d{2} мин."
        time_list = re.findall(time_regex, self.result_text)
        time_objs = [datetime.strptime(t, '%H час. %M мин.') for t in time_list]
        return [t.strftime('%H:%M') for t in time_objs]

    def get_form(self):
        form_regex = r"очной|заочной|очно-заочной"
        form_match = re.search(form_regex, self.result_text)
        return form_match.group(0)

    def get_type(self):
        type_regex = r"очередное|внеочередное"
        type_match = re.search(type_regex, self.result_text)
        return type_match.group(0)

    def get_questions(self):
        need_to_write = False
        questions = []
        for paragraph in self.doc.paragraphs:
            if "Решение:" in paragraph.text:
                need_to_write = True
            elif "Приложение:" in paragraph.text:
                need_to_write = False
            elif need_to_write:
                questions.append(paragraph.text)
        return questions

    def format_datetime(self, form, dates, times):
        if form == 'очной':
            return f'{times[0]}-{times[1]} {dates[0]}'
        elif form == 'заочной':
            return f'{times[0]}-{times[1]} {dates[0]}-{dates[1]}'
        else:
            return f'{times[0]}-{times[1]} {dates[0]} / {times[2]}-{times[3]} {dates[1]}-{dates[2]}'

    def get_place(self, form):
        if form == 'очной':
            return 'Ул. Ленина, д.14'
        elif form == 'заочной':
            return 'Сайт СНТ'
        else:
            return 'Ул. Ленина, д.14 / Сайт СНТ'

    def get_protocol_creating_date(self):
        paragraph = self.doc.paragraphs[0]
        date_regex = r"\d{2}\.\d{2}\.\d{4}"
        match = re.search(date_regex, paragraph.text)
        return match.group(0)
