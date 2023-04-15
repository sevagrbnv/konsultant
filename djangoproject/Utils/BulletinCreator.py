import docx
from datetime import datetime
import time as t
from docx.shared import Pt


def create_bulletin(name, type, date_time, form, questions):
    new_type = type.replace('ое', 'ого')
    time, date = date_time.split()
    start_time, end_time = time.split('-')
    start_time = start_time.replace(':', ' час. ') + ' мин.'
    end_time = end_time.replace(':', ' час. ') + ' мин.'
    date_time_end = f"{date} с {start_time} до {end_time}"
    doc = docx.Document(name)

    for para in doc.paragraphs:
        text = para.text
        text = text.replace('<<type>>', new_type)
        text = text.replace('<<date>>', date_time_end)
        text = text.replace('<<form>>', form)
        para.text = text

    for para in doc.paragraphs:
        text = para.text
        if "ПОВЕСТКА:" in para.text:
            for question in questions:
                doc.add_paragraph(question)
        para.text = text

    for question in range(len(questions)):
        doc.add_paragraph(f'{question + 1}. Вопрос повестки дня: ')
        text = questions[question][3:]
        doc.add_paragraph(text)
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'За'
        hdr_cells[1].text = 'Против'
        hdr_cells[2].text = 'Воздержался'

    for i in range(3):
        doc.add_paragraph()
    doc.add_paragraph('Дата ______________')
    doc.add_paragraph('Подпись ______________')
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    name = f'bulletin_{int(t.time())}.docx'
    doc.save(f'./uploads/{name}')
    return name