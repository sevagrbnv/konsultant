import docx
from datetime import datetime
import time as t
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


def create_notification(name, type, date_time, form, questions, protocol_date, current_date):
    new_type = type.replace('ое', 'ого')
    doc = docx.Document(name)
    for para in doc.paragraphs:
        if 'Место приема решений' in para.text:
            mark_paragraph = para
        text = para.text
        text = text.replace('<<type>>', new_type)
        text = text.replace('<<form>>', form)
        text = text.replace('<<current_date>>', current_date)
        para.text = text

    if form == 'очно-заочной':
        ochn_datetime, zaochn_datetime = date_time.split('/')
        och_time, och_date = ochn_datetime.split()
        och_start_time, och_end_time = och_time.split('-')
        och_start_time = och_start_time.replace(':', ' час. ') + ' мин.'
        och_end_time = och_end_time.replace(':', ' час. ') + ' мин.'
        zaoch_time, zaoch_dates = zaochn_datetime.split()
        zaoch_start_time, zaoch_end_time = zaoch_time.split('-')
        zaoch_start_time = zaoch_start_time.replace(':', ' час. ') + ' мин.'
        zaoch_end_time = zaoch_end_time.replace(':', ' час. ') + ' мин.'
        zaoch_start_date, zaoch_end_date = zaoch_dates.split('-')
        p = mark_paragraph.insert_paragraph_before(
            f'Дата, время начала и окончания проведения очной части общего собрания членов Товарищества СНТ 1:'
            f' Московская обл., Одинцовский р-н, д. Хлюпино, ул. Ленина д. 14, {och_date} c {och_start_time} до {och_end_time}.'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_paragraph.insert_paragraph_before(
            'Дата, время начала и окончания проведения заочной части с применением электронно-технических средств (на Официальном сайте http://www.snt1.ru.) общего собрания членов Товарищества СНТ 1.'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_paragraph.insert_paragraph_before(
            f'C {zaoch_start_date} c {zaoch_start_time} до {zaoch_end_time}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_paragraph.insert_paragraph_before(
            f'По {zaoch_end_date} c {zaoch_start_time} до {zaoch_end_time}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_paragraph.insert_paragraph_before(
            f'Перечень вопросов, подлежащих рассмотрению на общем собранию членов Товарищество {och_date}-{zaoch_end_date} (Протокол Правления от {protocol_date}).'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)

    elif form == 'очной':
        time, date = date_time.split()
        start_time, end_time = time.split('-')
        start_time = start_time.replace(':', ' час. ') + ' мин.'
        end_time = end_time.replace(':', ' час. ') + ' мин.'
        date_time_end = f"{date} с {start_time} до {end_time}"
        p = mark_paragraph.insert_paragraph_before('Дата, время начала и окончания проведения общего собрания членов Товарищества СНТ 1.', style='Normal')
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_paragraph.insert_paragraph_before(f'{date_time_end}')
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_paragraph.insert_paragraph_before('Место проведения общего собрания членов Товарищества: Московская обл., Одинцовский р-н, д. Хлюпино, Ул. Ленина д. 14.', style='Normal')
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_paragraph.insert_paragraph_before(f'Перечень вопросов, подлежащих рассмотрению на общем собрании членов Товарищества {date} (Протокол Правления от {protocol_date}):', style='Normal')
        p.paragraph_format.first_line_indent = Mm(7.5)

    else:
        pass

    for question in questions:
        p = mark_paragraph.insert_paragraph_before(question)
        p.paragraph_format.first_line_indent = Mm(7.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    doc.paragraphs[0].style = doc.styles['Normal']
    for para in range (2, len(doc.paragraphs)):
        doc.paragraphs[para].line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        doc.paragraphs[para].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p = doc.paragraphs[0]
    run = p.add_run('Уведомление членов СНТ о проведении собрания')
    run.font.size = Pt(10)
    name = f'notification_{int(t.time())}.docx'
    doc.save(f'./uploads/{name}')
    return name
