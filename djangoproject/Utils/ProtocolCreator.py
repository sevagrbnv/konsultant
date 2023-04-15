import docx
from datetime import datetime
import time as t
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


def create_protocol(name, type, date_time, form, questions, zaoch_list, protocol_date, current_date):
    new_form = form.replace('ой', 'ое')
    new_type = type.replace('ое', 'ого')
    new_type = new_type.capitalize()
    total_ochno = questions[0].yes + questions[0].no + questions[0].idk
    total_zaochno = zaoch_list[0][0]+zaoch_list[0][1]+zaoch_list[0][2]

    if (total_ochno + total_zaochno) > 0.5*49:
        quorum = 'да'
    else:
        quorum = 'нет'

    doc = docx.Document(name)
    for para in doc.paragraphs:
        if 'Повестка дня:' in para.text:
            mark_povestka = para
        text = para.text
        text = text.replace('<<current_date>>', current_date)
        text = text.replace('<<type>>', new_type)
        para.text = text




    if new_form == 'очно-заочное':
        ochn_datetime, zaochn_datetime = date_time.split('/')
        och_time, och_date = ochn_datetime.split()
        och_start_time, och_end_time = och_time.split('-')
        och_start_time = och_start_time.replace(':', ' час. ') + ' мин.'
        och_end_time = och_end_time.replace(':', ' час. ') + ' мин.'
        zaoch_time, zaoch_dates = zaochn_datetime.split()
        zaoch_start_time, zaoch_end_time = zaoch_time.split('-')
        zaoch_start_time = zaoch_start_time.replace(':', ' час. ') + ' мин.'
        zaoch_end_time = zaoch_end_time.replace(':', ' час. ') + ' мин.'
        zaoch_start_date, zaoch_end_date = zaoch_dates.split('-')
        p = mark_povestka.insert_paragraph_before(
            f'Место проведения: Московская обл., Одинцовский р-н, д. Хлюпино, Ул. Ленина д. 14.// Официальном сайте http://www.snt1.ru.'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Дата: с {och_date} по {zaoch_end_date}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Форма проведения: {new_form} собрание'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Дата, время, место проведения очного голосования: '
            f'Московская область, поселок Московский, ул. Ленина д. 15 {och_date} с {och_start_time} до {och_end_time}. '
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            'Дата, время, место проведения заочного голосования: '
            'Бюллетени принимались по адресу: Официальном сайте http://www.snt1.ru. '
            f'в период с {zaoch_start_date} по {zaoch_end_date} с {zaoch_start_time} до {zaoch_end_time}.'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Общее количество членов СНТ: 49'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Присутствовали: {total_ochno+total_zaochno}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Очно: {total_ochno}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Заочно: {total_zaochno}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Кворум [был или нет]: {quorum}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)

        i = 1
        for question in questions:
            p = doc.add_paragraph(question.question)
            p.paragraph_format.first_line_indent = Mm(7.5)
        for question in questions:
            p = doc.add_paragraph(f'{i} Вопрос повестки дня:')
            p.paragraph_format.first_line_indent = Mm(7.5)
            text = question.question[3:]
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Mm(7.5)
            p = doc.add_paragraph(f'Результаты голосования по {i}-му вопросу для очной части голосования общего собрания СНТ 1.')
            p.paragraph_format.first_line_indent = Mm(7.5)
            table = doc.add_table(rows=2, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            vote_cells = table.rows[1].cells
            hdr_cells[0].text = 'За'
            hdr_cells[1].text = 'Против'
            hdr_cells[2].text = 'Воздержался'
            vote_cells[0].text = f'{question.yes}'
            vote_cells[1].text = f'{question.no}'
            vote_cells[2].text = f'{question.idk}'
            p = doc.add_paragraph(
                f'Результаты голосования по {i}-му вопросу для заочной части голосования общего собрания СНТ 1.')
            p.paragraph_format.first_line_indent = Mm(7.5)
            table = doc.add_table(rows=2, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            vote_cells = table.rows[1].cells
            hdr_cells[0].text = 'За'
            hdr_cells[1].text = 'Против'
            hdr_cells[2].text = 'Воздержался'
            vote_cells[0].text = f'{zaoch_list[i-1][0]}'
            vote_cells[1].text = f'{zaoch_list[i-1][1]}'
            vote_cells[2].text = f'{zaoch_list[i-1][2]}'
            p = doc.add_paragraph(f'Принятое решение по {i}-му вопросу повестки для общего собрания:')
            p.paragraph_format.first_line_indent = Mm(7.5)
            p = doc.add_paragraph(f'{question.decision}')
            p.paragraph_format.first_line_indent = Mm(7.5)
            i += 1

        doc.add_paragraph()
        p = doc.add_paragraph('Решение принято на основании голосования в очно-заочной форме путем суммы голосов за очную часть голосования и заочную часть голосования.')
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(f'Общее количество членов СНТ 1: 49')
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Очно: {total_ochno}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Заочно: {total_zaochno}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        if quorum == 'да':
            p = mark_povestka.insert_paragraph_before(
                f'Кворум установлен ({total_ochno+total_zaochno} голосов в сумме)'
            )
            p.paragraph_format.first_line_indent = Mm(7.5)
        else:
            p = mark_povestka.insert_paragraph_before(
                f'Кворум не установлен ({total_ochno+total_zaochno} голосов в сумме)'
            )
            p.paragraph_format.first_line_indent = Mm(7.5)


    elif new_form == 'очное':
        time, date = date_time.split()
        start_time, end_time = time.split('-')
        start_time = start_time.replace(':', ' час. ') + ' мин.'
        end_time = end_time.replace(':', ' час. ') + ' мин.'
        date_time_end = f"{date} с {start_time} до {end_time}"
        p = mark_povestka.insert_paragraph_before(
            f'Место проведения: Московская обл., Одинцовский р-н, д. Хлюпино, Ул. Ленина д. 14.'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Дата: {date}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Форма проведения: {new_form} собрание'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Время проведения: Время начала собрания: {start_time}; Время окончания собрания: {end_time} '
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Общее количество членов СНТ: 49'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Присутствовали: {total_ochno}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)
        p = mark_povestka.insert_paragraph_before(
            f'Кворум [был или нет]: {quorum}'
        )
        p.paragraph_format.first_line_indent = Mm(7.5)

        i = 1
        for question in questions:
            p = doc.add_paragraph(question.question)
            p.paragraph_format.first_line_indent = Mm(7.5)

        for question in questions:
            p = doc.add_paragraph(f'{i} Вопрос повестки дня:')
            p.paragraph_format.first_line_indent = Mm(7.5)
            text = question.question[3:]
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Mm(7.5)
            p = doc.add_paragraph(f'Результаты голосования по {i}-му вопросу для общего собрания СНТ 1.')
            p.paragraph_format.first_line_indent = Mm(7.5)
            table = doc.add_table(rows=2, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            vote_cells = table.rows[1].cells
            hdr_cells[0].text = 'За'
            hdr_cells[1].text = 'Против'
            hdr_cells[2].text = 'Воздержался'
            vote_cells[0].text = f'{question.yes}'
            vote_cells[1].text = f'{question.no}'
            vote_cells[2].text = f'{question.idk}'
            p = doc.add_paragraph(f'Принятое решение по {i}-му вопросу повестки для общего собрания:')
            p.paragraph_format.first_line_indent = Mm(7.5)
            p = doc.add_paragraph(f'{question.decision}')
            p.paragraph_format.first_line_indent = Mm(7.5)
            i += 1

    else:
        pass

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph('Председатель_____________[подпись] Кузнецов С.М.')
    p.paragraph_format.first_line_indent = Mm(7.5)
    p = doc.add_paragraph('Секретарь_________________[подпись]  Соколова Е.С')
    p.paragraph_format.first_line_indent = Mm(7.5)
    p = doc.add_paragraph(f'Дата {current_date}')
    p.paragraph_format.first_line_indent = Mm(7.5)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    for para in range (4, len(doc.paragraphs)):
        doc.paragraphs[para].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        doc.paragraphs[para].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    name = f'protocol_sobraniya_{int(t.time())}.docx'
    doc.save(f'./uploads/{name}')
    return name
